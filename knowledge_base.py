"""Multi-collection ChromaDB knowledge base management."""

import json
import os
import uuid
from datetime import datetime
from io import BytesIO
from typing import Optional

import chromadb
import docx
from pypdf import PdfReader
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import (
    CHROMA_COLLECTION_NAME,
    CHROMA_PERSIST_DIR,
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    EMBEDDING_MODEL,
    KB_REGISTRY_DIR,
    KNOWLEDGE_DOCS_DIR,
    RETRIEVAL_K,
)


# ── Singleton Chroma client ───────────────────────────────────────────────────

_client: Optional[chromadb.PersistentClient] = None


def _get_client() -> chromadb.PersistentClient:
    global _client
    if _client is None:
        _client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
    return _client


def get_embeddings() -> OpenAIEmbeddings:
    return OpenAIEmbeddings(model=EMBEDDING_MODEL)


# ── Legacy single-collection helpers (backward compat) ───────────────────────

def get_vector_store() -> Chroma:
    return Chroma(
        client=_get_client(),
        collection_name=CHROMA_COLLECTION_NAME,
        embedding_function=get_embeddings(),
    )


def knowledge_base_exists() -> bool:
    try:
        store = get_vector_store()
        return store._collection.count() > 0
    except Exception:
        return False


def get_chunk_count() -> int:
    try:
        store = get_vector_store()
        return store._collection.count()
    except Exception:
        return 0


def ingest_knowledge_base(file_bytes: bytes, filename: str) -> int:
    text = _extract_text(file_bytes, filename)
    if not text.strip():
        return 0
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    docs = splitter.create_documents(
        texts=[text], metadatas=[{"source": filename}]
    )
    store = get_vector_store()
    store.add_documents(docs)
    return len(docs)


def ingest_from_directory() -> int:
    if not os.path.isdir(KNOWLEDGE_DOCS_DIR):
        return 0
    total = 0
    for fname in os.listdir(KNOWLEDGE_DOCS_DIR):
        fpath = os.path.join(KNOWLEDGE_DOCS_DIR, fname)
        if not os.path.isfile(fpath):
            continue
        if not fname.lower().endswith((".txt", ".docx", ".pdf")):
            continue
        with open(fpath, "rb") as f:
            file_bytes = f.read()
        total += ingest_knowledge_base(file_bytes, fname)
    return total


# ── Multi-KB registry ─────────────────────────────────────────────────────────

def _kb_meta_path(kb_id: str) -> str:
    kb_dir = os.path.join(KB_REGISTRY_DIR, kb_id)
    os.makedirs(kb_dir, exist_ok=True)
    return os.path.join(kb_dir, "metadata.json")


def _collection_name(kb_id: str) -> str:
    return f"kb_{kb_id}"


def create_kb(name: str, description: str = "") -> dict:
    """Create a new named KB. Returns metadata dict."""
    kb_id = str(uuid.uuid4())
    meta = {
        "kb_id": kb_id,
        "name": name,
        "description": description,
        "collection_name": _collection_name(kb_id),
        "created_at": datetime.utcnow().isoformat(),
        "chunk_count": 0,
        "keywords": [],
        "sources": [],
    }
    with open(_kb_meta_path(kb_id), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    # Eagerly create the ChromaDB collection
    _get_client().get_or_create_collection(_collection_name(kb_id))
    return meta


def get_kb_meta(kb_id: str) -> Optional[dict]:
    path = _kb_meta_path(kb_id)
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _save_kb_meta(meta: dict) -> None:
    path = _kb_meta_path(meta["kb_id"])
    with open(path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def list_kbs() -> list[dict]:
    """Return list of all KB metadata dicts."""
    result = []
    if not os.path.isdir(KB_REGISTRY_DIR):
        return result
    for kb_id in os.listdir(KB_REGISTRY_DIR):
        meta = get_kb_meta(kb_id)
        if meta:
            # Sync chunk count from ChromaDB
            try:
                coll = _get_client().get_collection(_collection_name(kb_id))
                meta["chunk_count"] = coll.count()
            except Exception:
                pass
            result.append(meta)
    return result


def ingest_into_kb(kb_id: str, file_bytes: bytes, filename: str, source_url: str = "") -> int:
    """Ingest a file into a named KB. Returns chunks added."""
    text = _extract_text(file_bytes, filename)
    if not text.strip():
        return 0

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    docs = splitter.create_documents(
        texts=[text],
        metadatas=[{"source": filename, "source_url": source_url, "kb_id": kb_id}],
    )

    store = Chroma(
        client=_get_client(),
        collection_name=_collection_name(kb_id),
        embedding_function=get_embeddings(),
    )
    store.add_documents(docs)

    # Update metadata
    meta = get_kb_meta(kb_id) or {}
    meta["chunk_count"] = meta.get("chunk_count", 0) + len(docs)
    sources = meta.get("sources", [])
    sources.append({"filename": filename, "url": source_url, "ingested_at": datetime.utcnow().isoformat()})
    meta["sources"] = sources
    _save_kb_meta(meta)

    return len(docs)


def search_kb(kb_id: str, query: str, k: int = RETRIEVAL_K) -> list[Document]:
    """Search a single KB by ID."""
    store = Chroma(
        client=_get_client(),
        collection_name=_collection_name(kb_id),
        embedding_function=get_embeddings(),
    )
    return store.similarity_search(query, k=k)


def search_across_kbs(kb_ids: list[str], query: str, k: int = RETRIEVAL_K) -> list[Document]:
    """Search across multiple KBs and merge results."""
    all_docs = []
    per_kb_k = max(2, k // max(len(kb_ids), 1))
    for kb_id in kb_ids:
        try:
            docs = search_kb(kb_id, query, k=per_kb_k)
            all_docs.extend(docs)
        except Exception:
            continue
    return all_docs[:k]


def get_retriever_for_kbs(kb_ids: list[str], k: int = RETRIEVAL_K):
    """Return a retriever that searches across the specified KBs."""
    from langchain_core.retrievers import BaseRetriever
    from langchain_core.callbacks import CallbackManagerForRetrieverRun

    class MultiKBRetriever(BaseRetriever):
        kb_ids_: list[str]
        k_: int

        def _get_relevant_documents(
            self, query: str, *, run_manager: CallbackManagerForRetrieverRun
        ) -> list[Document]:
            return search_across_kbs(self.kb_ids_, query, k=self.k_)

    return MultiKBRetriever(kb_ids_=kb_ids, k_=k)


def set_kb_keywords(kb_id: str, keywords: list[str]) -> bool:
    """Set enrichment keywords for a KB."""
    meta = get_kb_meta(kb_id)
    if meta is None:
        return False
    meta["keywords"] = keywords
    _save_kb_meta(meta)
    return True


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_text(file_bytes: bytes, filename: str) -> str:
    name_lower = filename.lower()
    if name_lower.endswith(".docx"):
        doc = docx.Document(BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n\n".join(parts)
    elif name_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    elif name_lower.endswith(".pdf"):
        reader = PdfReader(BytesIO(file_bytes))
        parts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n\n".join(parts)
    elif name_lower.endswith(".html") or name_lower.endswith(".htm"):
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(file_bytes, "html.parser")
            return soup.get_text(separator="\n")
        except ImportError:
            import re
            text = file_bytes.decode("utf-8", errors="replace")
            return re.sub(r"<[^>]+>", " ", text)
    elif name_lower.endswith(".csv"):
        import csv
        import io
        text = file_bytes.decode("utf-8", errors="replace")
        reader_csv = csv.reader(io.StringIO(text))
        rows = ["\t".join(row) for row in reader_csv]
        return "\n".join(rows)
    else:
        return ""
