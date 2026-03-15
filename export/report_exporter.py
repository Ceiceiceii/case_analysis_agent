"""Report export utilities: Markdown and DOCX."""

import io
from typing import Optional


def export_to_markdown(report_dict: dict) -> str:
    """Convert a report dict to a Markdown string."""
    if not report_dict:
        return ""

    title = report_dict.get("title", "Consulting Report")
    client = report_dict.get("client_name", "")
    version = report_dict.get("version", 1)
    status = report_dict.get("status", "draft")
    generated_at = report_dict.get("generated_at", "")
    content = report_dict.get("content", "")

    header = f"""# {title}

**Version:** {version} | **Status:** {status.upper()}
**Generated:** {generated_at}
{"**Client:** " + client if client else ""}

---

"""
    return header + content


def export_to_docx(report_dict: dict, filename: str = "report.docx") -> bytes:
    """
    Convert a report dict to a DOCX file.
    Returns bytes of the DOCX file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    doc = Document()

    title = report_dict.get("title", "Consulting Report")
    client = report_dict.get("client_name", "")
    version = report_dict.get("version", 1)
    status = report_dict.get("status", "draft")
    generated_at = report_dict.get("generated_at", "")
    content = report_dict.get("content", "")

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {version} | {status.upper()} | {generated_at}").italic = True
    if client:
        doc.add_paragraph(f"Client: {client}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("─" * 60)

    # Parse and add content
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(stripped.strip("*")).bold = True
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
