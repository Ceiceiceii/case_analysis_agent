from io import BytesIO

import docx
from pypdf import PdfReader


def parse_uploaded_file(uploaded_file) -> str:
    """Parse an uploaded file and return its text content."""
    name = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    if name.endswith(".docx"):
        return _parse_docx(file_bytes)
    elif name.endswith(".txt"):
        return _parse_txt(file_bytes)
    elif name.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    elif name.endswith(".html") or name.endswith(".htm"):
        return _parse_html(file_bytes)
    elif name.endswith(".csv"):
        return _parse_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {name}.")


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file, including paragraphs and tables."""
    doc = docx.Document(BytesIO(file_bytes))
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n\n".join(parts)


def _parse_txt(file_bytes: bytes) -> str:
    """Decode a plain text file from bytes."""
    return file_bytes.decode("utf-8")


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(BytesIO(file_bytes))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _parse_html(file_bytes: bytes) -> str:
    """Extract text from an HTML file."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(file_bytes, "html.parser")
        return soup.get_text(separator="\n")
    except ImportError:
        import re
        text = file_bytes.decode("utf-8", errors="replace")
        return re.sub(r"<[^>]+>", " ", text)


def _parse_csv(file_bytes: bytes) -> str:
    """Convert CSV content to tab-delimited text."""
    import csv
    import io
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    return "\n".join("\t".join(row) for row in reader)
